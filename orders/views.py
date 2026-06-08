import pypdf
import docx
import os
from django.utils import timezone
from django.core.signing import TimestampSigner, SignatureExpired, BadSignature
from django.core.mail import send_mail
from django.http import FileResponse, HttpResponse
from django.db.models import Q
from django.conf import settings
from rest_framework import status, permissions, generics
from rest_framework.views import APIView
from rest_framework.response import Response

from .models import Order, PricingConfig
from .serializers import OrderSerializer, PricingConfigSerializer
from accounts.models import User
from accounts.permissions import IsSuperAdmin, IsCollegeAdmin
from colleges.models import College

# Create signer for secure download link
signer = TimestampSigner()

def get_word_count(file):
    name = file.name.lower()
    text = ""
    try:
        # Seek file to start
        file.seek(0)
        if name.endswith('.pdf'):
            reader = pypdf.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif name.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        else:
            text = file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        print(f"Error parsing file for word count: {e}")
        text = "fallback text"
    
    words = text.split()
    count = len(words)
    return max(count, 1)


class WordCountEstimateView(APIView):
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        if 'file' not in request.FILES:
            return Response({"error": "No file uploaded"}, status=status.HTTP_400_BAD_REQUEST)
        
        uploaded_file = request.FILES['file']
        is_express = request.data.get('is_express') == 'true' or request.data.get('is_express') is True
        has_suggestions = request.data.get('has_suggestions') == 'true' or request.data.get('has_suggestions') is True
        
        word_count = get_word_count(uploaded_file)
        
        # Calculate pricing
        config = PricingConfig.objects.last()
        if not config:
            config = PricingConfig.objects.create(
                per_word_rate=0.50,
                express_fee=500.00,
                editing_suggestions_fee=299.00,
                referral_credit=100.00
            )
            
        base_price = word_count * config.per_word_rate
        total_price = base_price
        
        if is_express:
            total_price += config.express_fee
        if has_suggestions:
            total_price += config.editing_suggestions_fee
            
        return Response({
            "filename": uploaded_file.name,
            "word_count": word_count,
            "base_price": round(float(base_price), 2),
            "express_fee": round(float(config.express_fee), 2) if is_express else 0.0,
            "editing_suggestions_fee": round(float(config.editing_suggestions_fee), 2) if has_suggestions else 0.0,
            "total_price": round(float(total_price), 2)
        })


class OrderListCreateView(generics.ListCreateAPIView):
    serializer_class = OrderSerializer

    def get_queryset(self):
        user = self.request.user
        if user.role == 'super_admin':
            return Order.objects.all().order_by('-created_at')
        elif user.role == 'college_admin':
            # Submissions for all students in this college
            if not user.college:
                return Order.objects.none()
            
            # Apply filters
            queryset = Order.objects.filter(college=user.college)
            department = self.request.query_params.get('department')
            start_date = self.request.query_params.get('start_date')
            end_date = self.request.query_params.get('end_date')
            min_similarity = self.request.query_params.get('min_similarity')
            max_similarity = self.request.query_params.get('max_similarity')
            
            if department:
                queryset = queryset.filter(department__iexact=department)
            if start_date:
                queryset = queryset.filter(created_at__date__gte=start_date)
            if end_date:
                queryset = queryset.filter(created_at__date__lte=end_date)
            if min_similarity:
                queryset = queryset.filter(similarity_score__gte=float(min_similarity))
            if max_similarity:
                queryset = queryset.filter(similarity_score__lte=float(max_similarity))
                
            return queryset.order_by('-created_at')
        else:
            # Normal B2C user sees only their own orders
            return Order.objects.filter(user=user).order_by('-created_at')

    def create(self, request, *args, **kwargs):
        user = request.user
        if 'document' not in request.FILES:
            return Response({"error": "Document file is required"}, status=status.HTTP_400_BAD_REQUEST)
            
        file = request.FILES['document']
        is_express = request.data.get('is_express') == 'true' or request.data.get('is_express') is True
        has_suggestions = request.data.get('has_suggestions') == 'true' or request.data.get('has_suggestions') is True
        is_b2b_submission = request.data.get('is_b2b') == 'true' or request.data.get('is_b2b') is True

        word_count = get_word_count(file)
        config = PricingConfig.objects.last()
        if not config:
            config = PricingConfig.objects.create(
                per_word_rate=0.50,
                express_fee=500.00,
                editing_suggestions_fee=299.00,
                referral_credit=100.00
            )

        if is_b2b_submission:
            # Check B2B eligibility
            if not user.college:
                return Response({"error": "Your account is not associated with any college B2B credits."}, status=status.HTTP_400_BAD_REQUEST)
            if user.college.credits < 1:
                return Response({"error": "Insufficient credits remaining for your college account. Please contact college admin."}, status=status.HTTP_400_BAD_REQUEST)
            
            # Deduct 1 credit
            college = user.college
            college.credits -= 1
            college.save()

            # Check low credit warning (credits < 20% of allocated)
            if college.allocated_credits > 0 and (college.credits / college.allocated_credits) < 0.20:
                # Log warning
                print(f"[ALERT] College '{college.college_name}' credits are low! ({college.credits} remaining out of {college.allocated_credits})")
                # Send email to college admin if email exists
                if college.contact_email:
                    try:
                        send_mail(
                            'Low Credit Alert - Plagiarism Checker Platform',
                            f"Dear Admin,\n\nYour college account credits are running low. Remaining balance: {college.credits} out of {college.allocated_credits}.\n\nPlease renew/top up credits from the dashboard.\n\nBest regards,\nPlagiarism Checker Team",
                            'admin@plagiarismplatform.com',
                            [college.contact_email],
                            fail_silently=True
                        )
                    except Exception as e:
                        print(f"Failed to send low credit email: {e}")

            # Create Order
            order = Order.objects.create(
                user=user,
                document=file,
                word_count=word_count,
                price=0.00,
                status='Submitted',
                is_express=is_express,
                has_editing_suggestions=has_suggestions,
                is_b2b=True,
                college=college,
                department=user.department or 'General'
            )
            serializer = OrderSerializer(order)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
            
        else:
            # B2C Billed Order
            base_price = word_count * config.per_word_rate
            total_price = base_price
            if is_express:
                total_price += config.express_fee
            if has_suggestions:
                total_price += config.editing_suggestions_fee

            # Create Order (status: Submitted, will wait for Payment verification to proceed)
            # Or in this view we create order in 'Processing' or 'Submitted' state
            order = Order.objects.create(
                user=user,
                document=file,
                word_count=word_count,
                price=total_price,
                status='Submitted',
                is_express=is_express,
                has_editing_suggestions=has_suggestions,
                is_b2b=False
            )
            serializer = OrderSerializer(order)
            return Response(serializer.data, status=status.HTTP_201_CREATED)


class OrderDetailView(generics.RetrieveAPIView):
    queryset = Order.objects.all()
    serializer_class = OrderSerializer


class AddEditingSuggestionsView(APIView):
    def post(self, request, pk):
        try:
            order = Order.objects.get(pk=pk)
        except Order.DoesNotExist:
            return Response({"error": "Order not found"}, status=status.HTTP_404_NOT_FOUND)

        if order.has_editing_suggestions:
            return Response({"message": "Editing suggestions already added"}, status=status.HTTP_200_OK)

        config = PricingConfig.objects.last()
        upsell_fee = config.editing_suggestions_fee if config else 299.00
        
        order.has_editing_suggestions = True
        order.price += upsell_fee
        order.save()

        # If there is a payment record, update its amount as well
        if hasattr(order, 'payment'):
            payment = order.payment
            payment.amount += upsell_fee
            payment.save()

        return Response({
            "message": "Editing suggestions added successfully",
            "price": order.price
        })


class DownloadReportView(APIView):
    permission_classes = [permissions.AllowAny] # Anyone with the secure signed URL can download

    def get(self, request, pk):
        try:
            order = Order.objects.get(pk=pk)
        except Order.DoesNotExist:
            return Response({"error": "Order not found"}, status=status.HTTP_404_NOT_FOUND)

        token = request.query_params.get('token')
        if not token:
            return Response({"error": "Secure download token is missing"}, status=status.HTTP_400_BAD_REQUEST)

        # Verify Signature & Age
        try:
            # Max age of 48 hours (48 * 3600 = 172800 seconds)
            unsigned_id = signer.unsign(token, max_age=172800)
            if int(unsigned_id) != order.id:
                return Response({"error": "Invalid token signature"}, status=status.HTTP_400_BAD_REQUEST)
        except SignatureExpired:
            return Response({"error": "This report link has expired (valid for 48 hours only)"}, status=status.HTTP_400_BAD_REQUEST)
        except BadSignature:
            return Response({"error": "Invalid secure download link signature"}, status=status.HTTP_400_BAD_REQUEST)

        # Check report_uploaded_at as a secondary safety check
        if order.report_uploaded_at:
            delta = timezone.now() - order.report_uploaded_at
            if delta.total_seconds() > 172800:
                return Response({"error": "Report download window of 48 hours has expired"}, status=status.HTTP_400_BAD_REQUEST)

        if not order.report_file:
            return Response({"error": "Report file has not been uploaded yet"}, status=status.HTTP_400_BAD_REQUEST)

        # Serve file response
        file_path = order.report_file.path
        if not os.path.exists(file_path):
            return Response({"error": "Report file not found on server storage"}, status=status.HTTP_404_NOT_FOUND)

        return FileResponse(open(file_path, 'rb'), content_type='application/pdf')


class OrderInvoiceView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk):
        try:
            order = Order.objects.get(pk=pk)
        except Order.DoesNotExist:
            return Response({"error": "Order not found"}, status=status.HTTP_404_NOT_FOUND)

        # Check permissions
        if request.user.role != 'super_admin' and order.user != request.user:
            return Response({"error": "You do not have access to this invoice"}, status=status.HTTP_403_FORBIDDEN)

        # Import reportlab details
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from io import BytesIO

        try:
            # Use BytesIO instead of HttpResponse for safer PDF generation
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
            story = []

            styles = getSampleStyleSheet()
            
            # Styles
            title_style = ParagraphStyle(
                'TitleStyle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor("#1A202C"),
                spaceAfter=15
            )
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor("#4A5568"),
                spaceAfter=8
            )
            h2_style = ParagraphStyle(
                'H2Style',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor("#2D3748"),
                spaceAfter=10
            )

            # Invoice header and metadata
            story.append(Paragraph("TAX INVOICE", title_style))

            invoice_meta = [
                [Paragraph("<b>Invoice Number:</b>", body_style), Paragraph(f"INV-{order.id:06d}", body_style)],
                [Paragraph("<b>Invoice Date:</b>", body_style), Paragraph(order.created_at.strftime('%d %B %Y'), body_style)],
                [Paragraph("<b>Order Mode:</b>", body_style), Paragraph('B2B Credit' if order.is_b2b else 'B2C Checkout', body_style)],
                [Paragraph("<b>Order Status:</b>", body_style), Paragraph(order.status, body_style)],
            ]
            invoice_meta_table = Table(invoice_meta, colWidths=[120, 160], hAlign='RIGHT')
            invoice_meta_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                ('TEXTCOLOR', (0,0), (-1,-1), colors.HexColor("#4A5568")),
            ]))

            header_table = Table(
                [[Paragraph("", body_style), invoice_meta_table]],
                colWidths=[320, 230]
            )
            header_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(header_table)
            story.append(Spacer(1, 20))

            # Get user details safely
            user_full_name = order.user.get_full_name() or order.user.username
            user_phone = getattr(order.user, 'phone', 'N/A') or 'N/A'
            billing_data = [
                [Paragraph("<b>Billed To:</b>", body_style), Paragraph("<b>Service Provider:</b>", body_style)],
                [
                    Paragraph(
                        f"Name: {user_full_name}<br/>Email: {order.user.email}<br/>Phone: {user_phone}",
                        body_style
                    ),
                    Paragraph(
                        "Plagiarism Checker Platform Inc.<br/>Support: support@plagiarismplatform.com<br/>Web: plagiarismplatform.com",
                        body_style
                    )
                ]
            ]
            billing_table = Table(billing_data, colWidths=[275, 275])
            billing_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('LINEBEFORE', (1,0), (1,-1), 0.5, colors.HexColor('#CBD5E0')),
            ]))
            story.append(billing_table)
            story.append(Spacer(1, 25))

            story.append(Paragraph("Order Summary", h2_style))

            # Get document name safely
            doc_name = os.path.basename(order.document.name) if order.document else "Document"
            
            table_data = [
                [Paragraph("<b>Item</b>", body_style), Paragraph("<b>Description</b>", body_style), Paragraph("<b>Qty</b>", body_style), Paragraph("<b>Amount</b>", body_style)],
                [
                    Paragraph("Plagiarism Check", body_style),
                    Paragraph(doc_name, body_style),
                    Paragraph(f"{order.word_count} words", body_style),
                    Paragraph(f"₹ {order.price:.2f}", body_style)
                ]
            ]
            if order.is_express:
                table_data.append([
                    Paragraph("Express Check Premium", body_style),
                    Paragraph("Priority queue service", body_style),
                    Paragraph("1", body_style),
                    Paragraph("Included", body_style)
                ])
            if order.has_editing_suggestions:
                table_data.append([
                    Paragraph("Editing Suggestions Addon", body_style),
                    Paragraph("Grammar and phrasing guidance", body_style),
                    Paragraph("1", body_style),
                    Paragraph("Included", body_style)
                ])

            table_data.append([
                '',
                '',
                Paragraph("<b>Total Paid</b>", body_style),
                Paragraph(f"<b>₹ {order.price:.2f}</b>", body_style)
            ])

            summary_table = Table(table_data, colWidths=[170, 235, 80, 110])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#F7FAFC")),
                ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor("#2D3748")),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('ALIGN', (2,1), (3,-1), 'RIGHT'),
                ('LINEABOVE', (0,-1), (-1,-1), 1, colors.HexColor("#2D3748")),
                ('BOTTOMPADDING', (0,0), (-1,0), 10),
                ('BOTTOMPADDING', (0,1), (-1,-2), 6),
                ('BOTTOMPADDING', (2,-1), (3,-1), 10),
                ('GRID', (0,0), (-1,-2), 0.5, colors.HexColor("#E2E8F0")),
            ]))
            story.append(summary_table)
            story.append(Spacer(1, 30))

            story.append(Paragraph("Thank you for choosing our platform!", body_style))
            story.append(Paragraph("If you have any questions, please reach out to support@plagiarismplatform.com.", body_style))

            doc.build(story)
            
            # Get PDF bytes
            pdf_buffer.seek(0)
            pdf_bytes = pdf_buffer.getvalue()
            pdf_buffer.close()

            # Return PDF response
            response = HttpResponse(pdf_bytes, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="Invoice_{order.id}.pdf"'
            return response
            
        except Exception as e:
            print(f"Error generating invoice for order {pk}: {str(e)}")
            import traceback
            traceback.print_exc()
            return Response({"error": f"Failed to generate invoice: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class SuperAdminOrderQueueView(generics.ListAPIView):
    serializer_class = OrderSerializer
    permission_classes = [IsSuperAdmin]

    def get_queryset(self):
        # Return only active paid B2C orders or B2B orders, sorted with express first, excluding completed ones
        return Order.objects.filter(
            Q(is_b2b=True) | Q(payment__status='Paid') | Q(price=0.00)
        ).exclude(status='Report Ready').order_by('-is_express', '-created_at')


class SuperAdminUpdateOrderView(APIView):
    permission_classes = [IsSuperAdmin]

    def post(self, request, pk):
        try:
            order = Order.objects.get(pk=pk)
        except Order.DoesNotExist:
            return Response({"error": "Order not found"}, status=status.HTTP_404_NOT_FOUND)

        action = request.data.get('action') # 'start_processing' or 'complete'
        
        if action == 'start_processing':
            order.status = 'Processing'
            order.save()
            return Response({
                "message": "Order marked as Processing",
                "status": order.status
            })
            
        elif action == 'complete':
            similarity_score = request.data.get('similarity_score')
            report_file = request.FILES.get('report_file')
            
            if similarity_score is None:
                return Response({"error": "Similarity score is required"}, status=status.HTTP_400_BAD_REQUEST)
            if not report_file:
                return Response({"error": "Report PDF file is required"}, status=status.HTTP_400_BAD_REQUEST)

            try:
                similarity_score = float(similarity_score)
            except ValueError:
                return Response({"error": "Similarity score must be a number"}, status=status.HTTP_400_BAD_REQUEST)

            order.similarity_score = similarity_score
            order.report_file = report_file
            order.status = 'Report Ready'
            order.report_uploaded_at = timezone.now()
            order.save()

            # Generate expiring signed download URL
            token = signer.sign(str(order.id))
            # Create absolute URL or relative download link
            secure_link = f"/api/orders/{order.id}/download-report/?token={token}"

            # Simulate Notifications (Email & WhatsApp)
            print(f"[NOTIFICATION] Email and WhatsApp sent to user '{order.user.username}' (phone: {order.user.phone or 'N/A'}, email: {order.user.email})")
            print(f"[WHATSAPP SIMULATION] Message: Your plagiarism report for {os.path.basename(order.document.name)} is ready! Similarity: {similarity_score}%. Download link valid for 48 hours: {secure_link}")

            try:
                # Send Email
                send_mail(
                    'Your Plagiarism Check Report is Ready!',
                    f"Hi {order.user.username},\n\nYour document '{os.path.basename(order.document.name)}' has been verified.\nSimilarity Score: {similarity_score}%\n\nYou can access the secure download link below. Note: This link is valid only for 48 hours:\n{request.build_absolute_uri(secure_link)}\n\nBest regards,\nPlagiarism Checker Platform Support Team",
                    'support@plagiarismplatform.com',
                    [order.user.email],
                    fail_silently=True
                )
            except Exception as ex:
                print(f"Failed to send notification email: {ex}")

            return Response({
                "message": "Order marked as Complete. Notifications triggered.",
                "status": order.status,
                "similarity_score": order.similarity_score,
                "secure_download_url": secure_link
            })
            
        else:
            return Response({"error": "Invalid action. Must be 'start_processing' or 'complete'"}, status=status.HTTP_400_BAD_REQUEST)


class PricingConfigView(APIView):
    # Retrieve configuration (available to all logged in users)
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        config = PricingConfig.objects.last()
        if not config:
            config = PricingConfig.objects.create(
                per_word_rate=0.50,
                express_fee=500.00,
                editing_suggestions_fee=299.00,
                referral_credit=100.00
            )
        serializer = PricingConfigSerializer(config)
        return Response(serializer.data)

    def post(self, request):
        # Update configuration (superadmin only)
        if request.user.role != 'super_admin':
            return Response({"error": "Only super admins can modify pricing configuration"}, status=status.HTTP_403_FORBIDDEN)
            
        config = PricingConfig.objects.last()
        if not config:
            config = PricingConfig()
            
        serializer = PricingConfigSerializer(config, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
